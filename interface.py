import sqlite3
import os
from datetime import datetime

from geometry_package.rectangle import Rectangle
from geometry_package.triangle import Triangle
from geometry_package.trapezoid import Trapezoid

# WORD-экспорт
try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    WORD_AVAILABLE = True
except ImportError:
    WORD_AVAILABLE = False


# =====================================================
#                 DATABASE MANAGER
# =====================================================

class DatabaseManager:
    """Менеджер для работы с SQLite"""

    def __init__(self, db_name: str = None):
        # путь к БД — из ENV или default
        self.db_name = db_name or os.getenv("DB_PATH", "geometry_calculations.db")
        self._ensure_db()

    def _ensure_db(self):
        """Создаёт таблицу, если её нет — обязательно для Docker"""
        try:
            with sqlite3.connect(self.db_name) as conn:
                cursor = conn.cursor()
                cursor.execute("""
                    CREATE TABLE IF NOT EXISTS calculations (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        shape_type TEXT NOT NULL,
                        parameters TEXT NOT NULL,
                        area REAL NOT NULL,
                        circumscribed_radius REAL,
                        inscribed_radius REAL,
                        timestamp DATETIME DEFAULT CURRENT_TIMESTAMP
                    )
                """)
                conn.commit()
        except sqlite3.Error as e:
            raise ConnectionError(f"Ошибка инициализации БД: {e}")

    def save_calculation(self, shape_type, parameters, area, circ, insc):
        with sqlite3.connect(self.db_name) as conn:
            cursor = conn.cursor()
            cursor.execute("""
                INSERT INTO calculations (shape_type, parameters, area, circumscribed_radius, inscribed_radius)
                VALUES (?, ?, ?, ?, ?)
            """, (shape_type, parameters, area, circ, insc))
            conn.commit()

    def get_history(self):
        with sqlite3.connect(self.db_name) as conn:
            cursor = conn.cursor()
            cursor.execute("""
                SELECT id, shape_type, parameters, area,
                       circumscribed_radius, inscribed_radius,
                       strftime('%d.%m.%Y %H:%M', timestamp)
                FROM calculations ORDER BY timestamp DESC LIMIT 20
            """)
            return cursor.fetchall()


# =====================================================
#                 MAIN CONSOLE APPLICATION
# =====================================================

class GeometryConsoleApp:
    def __init__(self):
        self.db = DatabaseManager()
        self.current_calculation = None

    def clear_screen(self):
        os.system('cls' if os.name == 'nt' else 'clear')

    def show_menu(self):
        self.clear_screen()
        print("\nКАЛЬКУЛЯТОР ГЕОМЕТРИЧЕСКИХ ФИГУР")
        print("1. Рассчитать прямоугольник")
        print("2. Рассчитать треугольник")
        print("3. Рассчитать трапецию")
        print("4. Показать историю расчетов")
        print("5. Выход\n")

    def get_float(self, prompt):
        while True:
            try:
                v = float(input(prompt))
                if v <= 0:
                    print("Введите положительное число!")
                    continue
                return v
            except ValueError:
                print("Ошибка ввода!")

    def ask_save_word(self):
        if not WORD_AVAILABLE:
            return
        while True:
            ans = input("Сохранить отчёт в Word? (да/нет): ").lower()
            if ans in ("да", "д", "yes", "y"):
                self.save_to_word()
                return
            elif ans in ("нет", "н", "no", "n"):
                return

    # ---------------- RECTANGLE -----------------
    def calculate_rectangle(self):
        w = self.get_float("Введите ширину: ")
        h = self.get_float("Введите высоту: ")

        rect = Rectangle(w, h)
        print("\nРЕЗУЛЬТАТЫ:")
        print(f"Площадь: {rect.area:.4f}")
        print(f"R описанной: {rect.circumscribed_radius:.4f}")
        print(f"R вписанной: {rect.inscribed_radius if rect.inscribed_radius else 'нет'}")

        self.db.save_calculation(
            "Rectangle",
            f"ширина={w}, высота={h}",
            rect.area,
            rect.circumscribed_radius,
            rect.inscribed_radius
        )

        self.current_calculation = ("Rectangle", rect, f"ширина={w}, высота={h}")
        self.ask_save_word()

    # ---------------- TRIANGLE -----------------
    def is_right_triangle(self, a, b, c):
        s = sorted([a, b, c])
        return abs(s[0] ** 2 + s[1] ** 2 - s[2] ** 2) < 1e-6

    def calculate_triangle(self):
        a = self.get_float("Сторона A: ")
        b = self.get_float("Сторона B: ")
        c = self.get_float("Сторона C: ")

        tri = Triangle(a, b, c)

        print("\nРЕЗУЛЬТАТЫ:")
        print(f"Площадь: {tri.area:.4f}")
        print(f"R описанной: {tri.circumscribed_radius:.4f}")
        print(f"R вписанной: {tri.inscribed_radius:.4f}")
        print("Тип: прямоугольный" if self.is_right_triangle(a, b, c) else "Тип: непрямоугольный")

        self.db.save_calculation(
            "Triangle",
            f"{a},{b},{c}",
            tri.area,
            tri.circumscribed_radius,
            tri.inscribed_radius
        )

        self.current_calculation = ("Triangle", tri, f"{a},{b},{c}")
        self.ask_save_word()

    # ---------------- TRAPEZOID -----------------
    def calculate_trapezoid(self):
        b1 = self.get_float("Первое основание: ")
        b2 = self.get_float("Второе основание: ")
        h = self.get_float("Высота: ")

        trap = Trapezoid(b1, b2, h)

        print("\nРЕЗУЛЬТАТЫ:")
        print(f"Площадь: {trap.area:.4f}")
        print(f"Боковая сторона: {trap.side_length:.4f}")
        print(f"R описанной: {trap.circumscribed_radius or 'нет'}")
        print(f"R вписанной: {trap.inscribed_radius or 'нет'}")

        self.db.save_calculation(
            "Trapezoid",
            f"{b1},{b2}, h={h}",
            trap.area,
            trap.circumscribed_radius,
            trap.inscribed_radius
        )

        self.current_calculation = ("Trapezoid", trap, f"{b1},{b2}, h={h}")
        self.ask_save_word()

    # ---------------- HISTORY -----------------
    def show_history(self):
        rows = self.db.get_history()
        print("\nИСТОРИЯ РАСЧЁТОВ:\n")

        if not rows:
            print("История пуста")
            input("\nНажмите Enter...")
            return

        print(f"{'ID':<4} {'Тип':<12} {'Параметры':<25} {'Площадь':<10} {'Rоп':<8} {'Rвп':<8} {'Время':<16}")
        print("-" * 80)

        for r in rows:
            id_, t, p, area, r1, r2, tm = r
            print(f"{id_:<4} {t:<12} {p:<25} {area:<10.2f} "
                  f"{(r1 or '—'):<8} {(r2 or '—'):<8} {tm}")

        input("\nНажмите Enter...")

    # ---------------- WORD EXPORT -----------------
    def save_to_word(self):
        if not WORD_AVAILABLE:
            print("Модуль python-docx не установлен.")
            return

        if not self.current_calculation:
            print("Нет расчета для экспорта.")
            return

        shape_name, obj, params = self.current_calculation

        doc = Document()
        doc.add_heading("Отчёт о геометрическом расчёте", 0)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run(f"Дата: {datetime.now().strftime('%d.%m.%Y %H:%M:%S')}").italic = True

        doc.add_heading("1. Тип фигуры", level=1)
        doc.add_paragraph(shape_name)

        doc.add_heading("2. Параметры", level=1)
        doc.add_paragraph(params)

        doc.add_heading("3. Результаты", level=1)
        table = doc.add_table(rows=1, cols=2)
        row = table.rows[0].cells
        row[0].text = "Параметр"
        row[1].text = "Значение"

        data = [
            ("Площадь", f"{obj.area:.4f}"),
            ("R описанной", obj.circumscribed_radius or "нет"),
            ("R вписанной", obj.inscribed_radius or "нет"),
        ]

        for k, v in data:
            r = table.add_row().cells
            r[0].text = k
            r[1].text = str(v)

        name = f"report_{shape_name}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
        doc.save(name)

        print(f"Файл сохранён: {name}")

    # ---------------- MAIN LOOP -----------------

    def run(self):
        while True:
            self.show_menu()
            c = input("Выберите действие (1-5): ").strip()

            if c == "1": self.calculate_rectangle()
            elif c == "2": self.calculate_triangle()
            elif c == "3": self.calculate_trapezoid()
            elif c == "4": self.show_history()
            elif c == "5": break
            else: print("Ошибка выбора!")


def main():
    GeometryConsoleApp().run()


if __name__ == "__main__":
    main()
